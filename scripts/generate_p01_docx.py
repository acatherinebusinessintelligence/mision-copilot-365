# -*- coding: utf-8 -*-
"""Genera la plantilla oficial MCP-365-P01 en formato .docx real."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "planillas" / "MCP365_P01_Formato_analisis_correo.docx"
PURPLE = "4A3DA6"
DARK = "221E40"
YELLOW = "FFEC00"
LIGHT = "F4F3F9"
HINT_BG = "FFF8CC"
RULE_BG = "EFEAFB"
MUTED = "5A5A72"


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def set_cell_text(cell, text, *, bold=False, color=None, size=10, fill=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    if fill:
        shade_cell(cell, fill)
    # vertical-ish padding via cell margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "80")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def style_header_row(row, headers):
    for i, h in enumerate(headers):
        set_cell_text(row.cells[i], h, bold=True, color="FFFFFF", size=10, fill=PURPLE)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    style_header_row(table.rows[0], headers)
    for r_i, row_data in enumerate(rows):
        fill = LIGHT if r_i % 2 == 1 else None
        for c_i, val in enumerate(row_data):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, size=10, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_callout(doc, title, body, bg):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(title)
    set_run(r1, size=10, bold=True, color=DARK)
    r2 = p.add_run(" " + body)
    set_run(r2, size=10, bold=False, color=DARK)
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=13, bold=True, color=PURPLE)
    # underline feel via bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), PURPLE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # Banner
    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, DARK)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r = p1.add_run("Formato adjunto · Análisis de correo operativo")
    set_run(r, size=16, bold=True, color=YELLOW)
    p2 = cell.add_paragraph()
    r2 = p2.add_run("Misión Copilot 365 · ECCO · Universidad Sergio Arboleda · Capacitación empresarial")
    set_run(r2, size=9, color="D5D2E6")
    doc.add_paragraph()

    # Chips row as a small table
    chips = doc.add_table(rows=1, cols=2)
    set_cell_text(chips.rows[0].cells[0], "MCP-365-P01", bold=True, color="FFFFFF", size=9, fill=PURPLE, center=True)
    set_cell_text(chips.rows[0].cells[1], "Outlook + Copilot", bold=True, color=DARK, size=9, fill=YELLOW, center=True)
    chips.rows[0].cells[0].width = Cm(4)
    chips.rows[0].cells[1].width = Cm(4.5)
    doc.add_paragraph()

    # Meta fields
    meta = add_table(
        doc,
        ["Campo", "Dato"],
        [
            ["Participante", "________________________________"],
            ["Área / rol", "________________________________"],
            ["Correo", "________________________________"],
            ["Fecha", "________________"],
            ["App usada", "Outlook + Copilot"],
            ["Estado", "☐ Borrador    ☐ Validado    ☐ Entregado"],
        ],
        col_widths=[4, 12],
    )
    doc.add_paragraph()

    add_callout(
        doc,
        "Regla de oro:",
        "Copilot propone; tú validas. No inventes fechas, responsables ni cifras. Si falta un dato, escribe “No especificado”. Diferencia hechos e inferencias.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Cómo usarla:",
        "1) Descarga este archivo .docx real · 2) Ábrelo en Word · 3) Abre Copilot sobre ESTE documento · 4) Pega el prompt del reto y el correo · 5) Completa solo las 14 celdas de Hallazgo y Evidencia · 6) Guarda como MCP365_P01_Analisis_correo_completado.docx · 7) Validación humana y control de calidad los completa la persona.",
        HINT_BG,
    )
    add_callout(
        doc,
        "Para Copilot:",
        "llena SOLO la sección 2 (columnas Hallazgo y Evidencia). Conserva colores, tablas y diseño. No recrees el archivo. Trabaja sobre una COPIA de esta plantilla .docx.",
        HINT_BG,
    )
    add_callout(
        doc,
        "Para la persona:",
        "no pidas a Copilot que complete Validación humana ni Control de calidad. Esas secciones son solo humanas.",
        RULE_BG,
    )

    heading(doc, "1. Identificación del mensaje (opcional · puede quedar en blanco)")
    add_table(
        doc,
        ["Campo", "Dato", "Fuente / evidencia"],
        [
            ["De (remitente y cargo)", "", ""],
            ["Para / CC", "", ""],
            ["Asunto", "", ""],
            ["Fecha y hora de envío", "", ""],
            ["Código / activo / circuito", "", ""],
        ],
        col_widths=[5, 5.5, 5.5],
    )

    heading(doc, "2. Análisis estructurado · LLENAR CON COPILOT")
    add_para(
        doc,
        "Formato obligatorio: Elemento | Hallazgo | Evidencia en el correo",
        size=10,
        bold=True,
        color="36428C",
        space_after=4,
    )
    add_table(
        doc,
        ["Elemento", "Hallazgo", "Evidencia en el correo"],
        [
            ["1) Objetivo principal", "", ""],
            ["2) Hechos verificables", "", ""],
            ["3) Fechas y horas", "", ""],
            ["4) Responsables", "", ""],
            ["5) Compromisos", "", ""],
            ["6) Decisiones pendientes", "", ""],
            ["7) Preguntas sin resolver", "", ""],
        ],
        col_widths=[4.5, 5.75, 5.75],
    )

    heading(doc, "8. Validación humana · NO LLENAR CON COPILOT")
    add_table(
        doc,
        ["Hallazgo a validar", "¿Por qué requiere revisión humana?", "Estado"],
        [
            ["", "", "☐ Pendiente  ☐ Validado"],
            ["", "", "☐ Pendiente  ☐ Validado"],
            ["", "", "☐ Pendiente  ☐ Validado"],
        ],
        col_widths=[5.3, 6.2, 4.5],
    )
    add_para(
        doc,
        "La sección “Control de calidad” al final del documento también queda vacía: la completa el participante, no Copilot.",
        size=9,
        color=MUTED,
    )

    heading(doc, "9. Control de calidad (obligatorio) · NO LLENAR CON COPILOT")
    add_table(
        doc,
        ["Criterio", "Cumple", "Observación"],
        [
            ["Los hechos coinciden con la fuente original", "☐ Sí   ☐ No", ""],
            ["No hay datos inventados por Copilot", "☐ Sí   ☐ No", ""],
            ["Se marcaron vacíos como “No especificado”", "☐ Sí   ☐ No", ""],
            ["Se separaron hechos e inferencias", "☐ Sí   ☐ No", ""],
            ["Listo para uso operativo / entrega", "☐ Sí   ☐ No", ""],
        ],
        col_widths=[8, 3.5, 4.5],
    )

    heading(doc, "10. Firmas · NO LLENAR CON COPILOT")
    sig = doc.add_table(rows=2, cols=3)
    sig.style = "Table Grid"
    labels = [
        "Elaboró (participante)",
        "Revisó (par o líder)",
        "Aprobó (si aplica)",
    ]
    for i, lab in enumerate(labels):
        set_cell_text(sig.rows[0].cells[i], lab, bold=True, size=10)
        set_cell_text(sig.rows[1].cells[i], "\n\n_______________________\nNombre / fecha", size=10)

    doc.add_paragraph()
    add_para(
        doc,
        "Documento generado para la experiencia formativa Misión Copilot 365. Uso académico-operativo. No sustituye procedimientos oficiales de la organización. Plantilla oficial en formato Word (.docx) editable por Copilot.",
        size=9,
        color=MUTED,
    )

    doc.save(OUT)
    print(f"Created: {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
