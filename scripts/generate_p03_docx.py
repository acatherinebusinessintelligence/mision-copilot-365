# -*- coding: utf-8 -*-
"""Genera la plantilla oficial MCP-365-P03 · Matriz de compromisos (llenado)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "planillas" / "MCP365_P03_Matriz_compromisos_reunion.docx"
PURPLE = "4A3DA6"
DARK = "221E40"
YELLOW = "FFEC00"
LIGHT = "F4F3F9"
HINT_BG = "FFF8CC"
RULE_BG = "EFEAFB"


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def set_cell_text(cell, text, *, bold=False, color=None, size=10, fill=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    if fill:
        shade_cell(cell, fill)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "60")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=9, fill=PURPLE)
    for r_i, row_data in enumerate(rows):
        fill = LIGHT if r_i % 2 == 1 else None
        for c_i, val in enumerate(row_data):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, size=9, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_callout(doc, title, body, bg):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    set_run(p.add_run(title), size=10, bold=True, color=DARK)
    set_run(p.add_run(" " + body), size=10, color=DARK)
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), size=13, bold=True, color=PURPLE)


def add_empty_box(doc, hint: str):
    add_para(doc, hint, size=9, color="5A5A72", space_after=4)
    table = doc.add_table(rows=3, cols=1)
    for row in table.rows:
        set_cell_text(row.cells[0], "", size=10, fill=LIGHT)
    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.4)
        s.bottom_margin = Cm(1.4)
        s.left_margin = Cm(1.6)
        s.right_margin = Cm(1.6)

    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, DARK)
    cell.text = ""
    set_run(cell.paragraphs[0].add_run("Plantilla de llenado · Matriz de compromisos post-reunión"), size=15, bold=True, color=YELLOW)
    p2 = cell.add_paragraph()
    set_run(p2.add_run("Misión Copilot 365 · ECCO · Universidad Sergio Arboleda"), size=9, color="D5D2E6")
    doc.add_paragraph()

    chips = doc.add_table(rows=1, cols=2)
    set_cell_text(chips.rows[0].cells[0], "MCP-365-P03", bold=True, color="FFFFFF", size=9, fill=PURPLE, center=True)
    set_cell_text(chips.rows[0].cells[1], "Teams + Word + Copilot", bold=True, color=DARK, size=9, fill=YELLOW, center=True)
    doc.add_paragraph()

    add_table(
        doc,
        ["Campo", "Dato"],
        [
            ["Participante", "________________________________"],
            ["Área / rol", "________________________________"],
            ["Correo", "________________________________"],
            ["Fecha", "________________"],
            ["App usada", "Teams + Word + Copilot"],
            ["Fuente usada", "MCP365_P03_Fuente_transcripcion_Teams_Proyecto_Horizonte.docx"],
            ["Estado", "☐ Borrador    ☐ Validado    ☐ Entregado"],
        ],
        col_widths=[4, 12],
    )
    doc.add_paragraph()

    add_callout(
        doc,
        "Archivos del reto:",
        "1) Fuente: MCP365_P03_Fuente_transcripcion_Teams_Proyecto_Horizonte.docx · "
        "2) Esta plantilla de llenado. Completa solo secciones 1–5. "
        "Propuestas no aprobadas van DENTRO de la sección 2 (no en otro archivo). "
        "Validado por, control de calidad y firmas: solo la persona.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Entrega:",
        "Guarda como MCP365_P03_Matriz_compromisos_completada.docx. Conserva tablas y diseño.",
        HINT_BG,
    )

    heading(doc, "1. Datos de la reunión · LLENAR CON COPILOT")
    add_table(
        doc,
        ["Campo", "Información"],
        [
            ["Nombre / código de reunión", ""],
            ["Fecha y hora", ""],
            ["Facilitador", ""],
            ["Participantes", ""],
            ["Objetivo declarado", ""],
        ],
        col_widths=[5, 11],
    )

    heading(doc, "2. Clasificación de la conversación · LLENAR CON COPILOT")
    add_para(
        doc,
        "Incluye en «Propuestas no aprobadas» una lista numerada (propuesta, quien la presentó, motivo, condición, evidencia). No generar archivo separado.",
        size=9,
        color="5A5A72",
        space_after=4,
    )
    add_table(
        doc,
        ["Tipo", "Contenido (con evidencia)", "Estado"],
        [
            ["Decisiones confirmadas", "", "☐ Cerrada"],
            ["Propuestas no aprobadas", "", "☐ Abierta"],
            ["Opiniones / preocupaciones", "", "☐ Registrada"],
            ["Preguntas sin resolver", "", "☐ Pendiente"],
        ],
        col_widths=[4, 9, 3],
    )

    heading(doc, "3. Matriz de compromisos · LLENAR CON COPILOT")
    add_para(
        doc,
        "Ocho filas. Columna «Validado por» permanece vacía. Filas no usadas: vacías (no «No especificado»).",
        size=9,
        color="5A5A72",
        space_after=4,
    )
    empty8 = [
        [str(i), "", "", "", "", "☐ Pendiente ☐ En curso ☐ Hecho", "", ""]
        for i in range(1, 9)
    ]
    add_table(
        doc,
        ["#", "Actividad", "Responsable", "Fecha límite", "Dependencia", "Estado", "Evidencia / entregable", "Validado por"],
        empty8,
        col_widths=[0.8, 2.5, 2, 1.8, 2, 2.4, 2.8, 1.7],
    )

    heading(doc, "4. Próxima agenda · LLENAR CON COPILOT")
    add_table(
        doc,
        ["Tema", "Objetivo", "Preparación requerida", "Dueño"],
        [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
        col_widths=[4, 4, 4, 4],
    )

    heading(doc, "5. Resumen para difusión (máx. 5 líneas) · LLENAR CON COPILOT")
    add_empty_box(doc, "Propósito, decisiones confirmadas, compromisos relevantes, pendientes y próxima reunión si está confirmada.")

    heading(doc, "6. Control de calidad · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Criterio", "Cumple", "Observaciones"],
        [
            ["Decisiones solo con confirmación explícita", "☐ Sí ☐ No", ""],
            ["Propuestas no aprobadas en sección 2 (no archivo aparte)", "☐ Sí ☐ No", ""],
            ["Columna Validado por vacía", "☐ Sí ☐ No", ""],
            ["Diseño de plantilla conservado", "☐ Sí ☐ No", ""],
        ],
        col_widths=[7, 3, 6],
    )

    heading(doc, "7. Firmas · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Rol", "Nombre", "Fecha"],
        [["Elaboró", "", ""], ["Revisó", "", ""], ["Aprobó", "", ""]],
        col_widths=[4, 8, 4],
    )

    add_para(
        doc,
        "Misión Copilot 365 · Formato oficial MCP-365-P03 · No reconstruir el diseño.",
        size=8,
        color="5A5A72",
        space_after=0,
    )

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
