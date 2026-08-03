# -*- coding: utf-8 -*-
"""Genera dos .docx del Reto 2:
1) Fuente de correos ST-Urb-03
2) Plantilla de llenado (sin correos embebidos)
"""
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from app import get_reto_r2_parts  # noqa: E402

OUT_FUENTE = ROOT / "planillas" / "MCP365_P02_Fuente_correos_ST-Urb-03.docx"
OUT_PLANILLA = ROOT / "planillas" / "MCP365_P02_Formato_cadena_correos.docx"
PURPLE = "4A3DA6"
DARK = "221E40"
YELLOW = "FFEC00"
LIGHT = "F4F3F9"
HINT_BG = "FFF8CC"
RULE_BG = "EFEAFB"
EMAIL_BG = "F7F9FC"


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def set_cell_text(cell, text, *, bold=False, color=None, size=10, fill=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    if fill:
        shade_cell(cell, fill)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "80")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def style_header_row(row, headers):
    for i, h in enumerate(headers):
        set_cell_text(row.cells[i], h, bold=True, color="FFFFFF", size=10, fill=PURPLE)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    style_header_row(table.rows[0], headers)
    for r_i, row_data in enumerate(rows):
        fill = LIGHT if r_i % 2 == 1 else None
        for c_i, val in enumerate(row_data):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, size=10, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_callout(doc, title, body, bg):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(title)
    set_run(r1, size=10, bold=True, color=DARK)
    r2 = p.add_run(" " + body)
    set_run(r2, size=10, bold=False, color=DARK)
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=13, bold=True, color=PURPLE)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), PURPLE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def add_banner(doc, title: str):
    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, DARK)
    cell.text = ""
    p1 = cell.paragraphs[0]
    set_run(p1.add_run(title), size=16, bold=True, color=YELLOW)
    p2 = cell.add_paragraph()
    set_run(
        p2.add_run("Misión Copilot 365 · ECCO · Universidad Sergio Arboleda"),
        size=9,
        color="D5D2E6",
    )
    doc.add_paragraph()


def add_chips(doc, left: str, right: str):
    chips = doc.add_table(rows=1, cols=2)
    set_cell_text(chips.rows[0].cells[0], left, bold=True, color="FFFFFF", size=9, fill=PURPLE, center=True)
    set_cell_text(chips.rows[0].cells[1], right, bold=True, color=DARK, size=9, fill=YELLOW, center=True)
    doc.add_paragraph()


def add_email_block(doc, index: int, part: dict):
    body = part["body_tpl"].format(sig=part["sig"]).strip()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, EMAIL_BG)
    cell.text = ""
    title = cell.paragraphs[0]
    set_run(title.add_run(f"Correo {index} · {part['date']}"), size=11, bold=True, color=DARK)

    def add_line(label: str, value: str):
        p = cell.add_paragraph()
        set_run(p.add_run(f"{label}: "), size=10, bold=True, color=PURPLE)
        set_run(p.add_run(value), size=10, bold=False, color=DARK)

    add_line("De", part["from_name"])
    add_line("Para", part["to_line"])
    add_line("Asunto", part["subject"])
    add_line("Fecha", part["date"])
    cell.add_paragraph()
    for para in body.split("\n"):
        p = cell.add_paragraph()
        set_run(p.add_run(para if para else " "), size=10, color=DARK)
    doc.add_paragraph()


def add_empty_box(doc, hint: str):
    add_para(doc, hint, size=9, color="5A5A72", space_after=4)
    table = doc.add_table(rows=3, cols=1)
    for row in table.rows:
        set_cell_text(row.cells[0], "", size=10, fill=LIGHT)
        row.cells[0].width = Cm(16)
    doc.add_paragraph()


def build_fuente():
    doc = Document()
    setup_page(doc)
    add_banner(doc, "Fuente documental · Cadena de correos ST-Urb-03")
    add_chips(doc, "MCP-365-P02-F", "Solo lectura · 4 correos")
    add_callout(
        doc,
        "Uso:",
        "Este archivo es la FUENTE. No es la plantilla de llenado. "
        "Ábrelo junto con MCP365_P02_Formato_cadena_correos.docx. "
        "Copilot analiza estos correos (De, Para, Asunto, Fecha, cuerpo) "
        "y completa solo la plantilla de llenado.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Regla de evidencia:",
        "Al citar, usa siempre: fecha + De + Asunto. "
        "Ejemplo: «04/03/2026 · Julián Pardo · Logística de Materiales · RE: adelanto de repuestos · ST-Urb-03».",
        HINT_BG,
    )
    heading(doc, "Correos de la cadena (no editar)")
    for i, part in enumerate(get_reto_r2_parts(), 1):
        add_email_block(doc, i, part)
    OUT_FUENTE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FUENTE)
    print(f"OK -> {OUT_FUENTE}")


def build_planilla():
    doc = Document()
    setup_page(doc)
    add_banner(doc, "Plantilla de llenado · Cadena de correos · ST-Urb-03")
    add_chips(doc, "MCP-365-P02", "Word + Copilot · Llenar")

    add_table(
        doc,
        ["Campo", "Dato"],
        [
            ["Participante", "________________________________"],
            ["Área / rol", "________________________________"],
            ["Fecha", "________________"],
            ["App usada", "Word + Copilot"],
            ["Fuente usada", "MCP365_P02_Fuente_correos_ST-Urb-03.docx"],
            ["Estado", "☐ Borrador    ☐ Validado    ☐ Entregado"],
        ],
        col_widths=[4, 12],
    )
    doc.add_paragraph()

    add_callout(
        doc,
        "Archivos del reto:",
        "1) MCP365_P02_Fuente_correos_ST-Urb-03.docx = correos (De/Para/Asunto). "
        "2) ESTE archivo = plantilla de llenado. Abre ambos en Word. "
        "Copilot debe leer la fuente y completar SOLO las celdas/cuadros vacíos de esta plantilla.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Para Copilot:",
        "No recrees el documento. No cambies títulos, tablas, colores ni diseño. "
        "No completes Validación humana ni Control de calidad. "
        "Si falta un dato en la fuente: escribe «No especificado».",
        HINT_BG,
    )
    add_callout(
        doc,
        "Entrega:",
        "Guarda ESTE archivo (ya lleno) como MCP365_P02_Cadena_correos_completado.docx.",
        HINT_BG,
    )

    heading(doc, "1. Cronología de la cadena · LLENAR CON COPILOT")
    add_para(
        doc,
        "Una fila por correo de la fuente. En «Evidencia» cita fecha + De + Asunto + fragmento textual.",
        size=9,
        color="5A5A72",
        space_after=4,
    )
    add_table(
        doc,
        ["#", "Fecha", "De (remitente)", "Asunto", "Hecho o cambio", "Compromiso / pendiente", "Evidencia (cita)"],
        [["1", "", "", "", "", "", ""],
         ["2", "", "", "", "", "", ""],
         ["3", "", "", "", "", "", ""],
         ["4", "", "", "", "", "", ""]],
        col_widths=[1, 2, 3, 3, 2.5, 2.5, 2.5],
    )

    heading(doc, "2. Cambios respecto al plan inicial · LLENAR CON COPILOT")
    add_table(
        doc,
        ["Tema", "Plan inicial (correo 01 mar)", "Plan vigente", "Quién lo indicó (De + Asunto)", "Pendiente de comunicar"],
        [
            ["Fecha / ventana", "", "", "", "☐ Sí ☐ No"],
            ["Horario / cierre de área", "", "", "", "☐ Sí ☐ No"],
            ["Personal / seguridad", "", "", "", "☐ Sí ☐ No"],
            ["Comunicación a usuarios", "", "", "", "☐ Sí ☐ No"],
        ],
        col_widths=[3, 3.5, 3, 4, 2.5],
    )

    heading(doc, "3. Compromisos vigentes · LLENAR CON COPILOT")
    add_table(
        doc,
        ["#", "Compromiso", "Responsable", "Origen (De + Asunto + fecha)", "Estado"],
        [
            ["1", "", "", "", "☐ Pendiente ☐ Hecho"],
            ["2", "", "", "", "☐ Pendiente ☐ Hecho"],
            ["3", "", "", "", "☐ Pendiente ☐ Hecho"],
            ["4", "", "", "", "☐ Pendiente ☐ Hecho"],
        ],
        col_widths=[1, 4.5, 3, 5, 2.5],
    )

    heading(doc, "4. Respuesta técnica · LLENAR CON COPILOT")
    add_empty_box(doc, "Precisión operativa: procedimiento, tiempos, controles, responsable. Solo con datos de la fuente.")

    heading(doc, "5. Respuesta ejecutiva (máx. 8 líneas) · LLENAR CON COPILOT")
    add_empty_box(doc, "Decisión, impacto, riesgo, siguiente paso y dueño. Cita el correo de aprobación (De + Asunto).")

    heading(doc, "6. Comunicación a usuarios / comunidad (máx. 6 líneas) · LLENAR CON COPILOT")
    add_empty_box(doc, "Lenguaje simple: qué ocurre, cuándo, cómo afecta, canal. Sin jerga. Basado en fechas vigentes de la cadena.")

    heading(doc, "7. Validación humana · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Pregunta", "Respuesta"],
        [
            ["¿La cronología cita De y Asunto de cada correo?", "☐ Sí ☐ No · Notas: ________"],
            ["¿Hay datos inventados?", "☐ No ☐ Sí (corregir): ________"],
            ["Validador / fecha", "________________"],
        ],
        col_widths=[7, 9],
    )

    heading(doc, "8. Control de calidad · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Criterio", "Cumple"],
        [
            ["Diseño de la plantilla conservado (sin reconstruir)", "☐ Sí ☐ No"],
            ["Tres audiencias completas y diferenciadas", "☐ Sí ☐ No"],
            ["Archivo guardado como MCP365_P02_Cadena_correos_completado.docx", "☐ Sí ☐ No"],
        ],
        col_widths=[12, 4],
    )

    OUT_PLANILLA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PLANILLA)
    print(f"OK -> {OUT_PLANILLA}")


def main():
    build_fuente()
    build_planilla()


if __name__ == "__main__":
    main()
