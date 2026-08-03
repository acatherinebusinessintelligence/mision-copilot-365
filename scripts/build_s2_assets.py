# -*- coding: utf-8 -*-
"""Genera los 6 activos de práctica MCP-365-S2 · Proyecto Horizonte (ficticio)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "planillas"

# Paleta corporativa
DARK = "211D40"
PURPLE = "4B3FAA"
YELLOW = "FFEC00"
LIGHT = "F4F3F9"
MUTED = "5A5A72"
WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# Helpers docx
# ---------------------------------------------------------------------------
def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(
    doc,
    text,
    *,
    size=11,
    bold=False,
    color=None,
    space_after=6,
    space_before=0,
    align=None,
):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    return p


def set_cell_text(
    cell,
    text,
    *,
    bold=False,
    color=None,
    size=10,
    fill=None,
    center=False,
):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    if fill:
        shade_cell(cell, fill)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "80")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_banner(doc, title: str, subtitle: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, DARK)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    set_run(r, size=16, bold=True, color=YELLOW)
    if subtitle:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(subtitle)
        set_run(r2, size=10, bold=False, color=WHITE)
    doc.add_paragraph()


def empty_box(doc, hint: str = "Escribir aquí…") -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, LIGHT)
    set_cell_text(cell, hint, color=MUTED, size=10)
    # altura visual mínima
    for _ in range(2):
        cell.add_paragraph()
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# 1. Correo HTML
# ---------------------------------------------------------------------------
def build_email_html(path: Path) -> None:
    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>NUEVO PROYECTO · HORIZONTE</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: Calibri, Segoe UI, Arial, sans-serif;
    background: #E8E6F0;
    color: #{DARK};
    padding: 24px 12px;
    line-height: 1.5;
  }}
  .shell {{
    max-width: 640px;
    margin: 0 auto;
    background: #fff;
    border: 1px solid #D0CCE0;
    box-shadow: 0 4px 18px rgba(33, 29, 64, 0.12);
  }}
  .bar {{
    background: #{DARK};
    height: 6px;
  }}
  .header {{
    background: #{PURPLE};
    color: #fff;
    padding: 18px 24px;
  }}
  .header .brand {{
    font-size: 12px;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: #{YELLOW};
    margin-bottom: 6px;
  }}
  .header h1 {{
    font-size: 18px;
    font-weight: 700;
  }}
  .meta {{
    background: #{LIGHT};
    padding: 14px 24px;
    border-bottom: 1px solid #D0CCE0;
    font-size: 13px;
  }}
  .meta dt {{
    float: left;
    width: 90px;
    font-weight: 700;
    color: #{PURPLE};
    clear: left;
  }}
  .meta dd {{
    margin-left: 100px;
    margin-bottom: 6px;
  }}
  .body {{
    padding: 22px 24px 28px;
    font-size: 14px;
  }}
  .body p {{ margin-bottom: 12px; }}
  .accent {{
    border-left: 4px solid #{YELLOW};
    background: #{LIGHT};
    padding: 12px 14px;
    margin: 16px 0;
  }}
  .attach {{
    display: flex;
    align-items: center;
    gap: 10px;
    border: 1px dashed #{PURPLE};
    padding: 10px 12px;
    margin: 16px 0;
    font-size: 13px;
  }}
  .attach .icon {{
    width: 36px;
    height: 36px;
    background: #{PURPLE};
    color: #{YELLOW};
    display: flex;
    align-items: center;
    justify-content: center;
    font-weight: 700;
    font-size: 11px;
    flex-shrink: 0;
  }}
  .sign {{
    margin-top: 20px;
    padding-top: 14px;
    border-top: 1px solid #D0CCE0;
    font-size: 13px;
    color: #{MUTED};
  }}
  .footer {{
    background: #{DARK};
    color: #C8C4D8;
    font-size: 11px;
    padding: 12px 24px;
    text-align: center;
  }}
  .note {{
    margin-top: 18px;
    font-size: 11px;
    color: #{MUTED};
    font-style: italic;
  }}
</style>
</head>
<body>
  <div class="shell">
    <div class="bar"></div>
    <div class="header">
      <div class="brand">MCP-365 · Simulación de correo · Caso ficticio</div>
      <h1>Autorización de inicio · Proyecto Horizonte</h1>
    </div>
    <dl class="meta">
      <dt>De:</dt>
      <dd>Dirección de Operaciones &lt;operaciones@ejemplo-corporativo.local&gt;</dd>
      <dt>Para:</dt>
      <dd>Oficina de PMO &lt;pmo@ejemplo-corporativo.local&gt;</dd>
      <dt>Asunto:</dt>
      <dd><strong>NUEVO PROYECTO · HORIZONTE</strong></dd>
      <dt>Fecha:</dt>
      <dd>3 de agosto de 2026 · 09:15 (hora local ficticia)</dd>
    </dl>
    <div class="body">
      <p>Estimado equipo de la Oficina de PMO,</p>
      <p>
        Por medio del presente se <strong>autoriza el inicio formal</strong> del
        <strong>Proyecto Horizonte</strong> (código preliminar
        <strong>PRJ-HZ-2026-01</strong>), caso de práctica MCP-365-S2.
      </p>
      <div class="accent">
        La Dirección de Operaciones solicita registrar el proyecto, validar la
        ficha adjunta y activar la cadena de automatización de la sesión
        (correo → adjunto → control Excel → informe → aprobación humana → PPT).
      </div>
      <p>
        Adjunto encontrarán la ficha de proyecto
        <strong>MCP365_S2_Ficha_Proyecto_Horizonte.docx</strong>, con objetivo,
        alcance preliminar, fechas, presupuesto de referencia y riesgos iniciales.
        Algunos campos figuran como <em>No especificado</em> a propósito, para
        ejercitar la extracción y el contraste con la validación humana.
      </p>
      <div class="attach">
        <div class="icon">DOC</div>
        <div>
          <strong>MCP365_S2_Ficha_Proyecto_Horizonte.docx</strong><br>
          Ficha de inicio · Proyecto Horizonte (ficticio)
        </div>
      </div>
      <p>
        Quedan a disposición para aclaraciones. No se incluyen credenciales reales
        ni datos de producción; todo el contenido es material didáctico.
      </p>
      <div class="sign">
        Cordialmente,<br>
        <strong style="color:#{DARK}">Dirección de Operaciones</strong><br>
        Organización de práctica MCP-365<br>
        Contacto: No especificado
      </div>
      <p class="note">
        Simulación offline para sesión MCP-365-S2. No enviar a buzones reales.
      </p>
    </div>
    <div class="footer">
      Paleta: #{DARK} · #{PURPLE} · #{YELLOW} &nbsp;|&nbsp; Proyecto Horizonte · Caso ficticio
    </div>
  </div>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


# ---------------------------------------------------------------------------
# 2. Ficha Word
# ---------------------------------------------------------------------------
def build_ficha_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    add_banner(
        doc,
        "MCP-365-S2 · Ficha Proyecto Horizonte",
        "Caso ficticio de práctica · Sin credenciales reales",
    )
    add_para(
        doc,
        "Formulario de inicio de proyecto. Completar / contrastar en la sesión.",
        size=10,
        color=MUTED,
        space_after=10,
    )

    fields = [
        ("Nombre del proyecto", "Proyecto Horizonte"),
        ("Código", "PRJ-HZ-2026-01"),
        (
            "Objetivo",
            "Habilitar un flujo piloto de gobernanza documental y control "
            "presupuestario para iniciativas de operaciones, con trazabilidad "
            "desde el correo de autorización hasta el informe aprobado.",
        ),
        (
            "Alcance",
            "Incluye: registro de recepción, ficha de proyecto, control de "
            "presupuesto y riesgos, informe ejecutivo y presentación post-aprobación. "
            "Excluye: despliegue productivos, integración con ERP real y datos personales.",
        ),
        ("Patrocinador", "Dirección de Operaciones"),
        ("Líder de proyecto", "No especificado"),
        ("Fecha de inicio", "2026-08-10"),
        ("Fecha de fin (plan)", "2026-11-30"),
        ("Presupuesto preliminar", "100 u.m."),
        (
            "Entregables",
            "1) Ficha de proyecto validada; 2) Libro de control Excel; "
            "3) Informe ejecutivo; 4) Presentación de cierre de fase; "
            "5) Evidencia de aprobación humana.",
        ),
        (
            "Dependencias",
            "Disponibilidad de plantillas MCP-365-S2; acceso al buzón de práctica "
            "de la PMO; definición del umbral de aprobación presupuestaria.",
        ),
        (
            "Restricciones",
            "Presupuesto tope 100 u.m.; duración máxima del piloto 16 semanas; "
            "sin uso de datos reales de clientes; aprobaciones humanas obligatorias.",
        ),
        (
            "Riesgos iniciales",
            "R1: Campos incompletos en la ficha (p. ej. líder No especificado) "
            "retrasan el arranque. R2: Desviación presupuestaria por cifras estimadas. "
            "R3: Confusión entre automatización y aprobación humana.",
        ),
        (
            "Criterios de éxito",
            "Cadena de 6 retos operativa; presupuesto y riesgos trazables; "
            "informe aprobado por humano (nunca por Copilot); presentación generada "
            "solo tras aprobación.",
        ),
        ("Observaciones / notas", "No especificado"),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        set_cell_text(row.cells[0], label, bold=True, color=WHITE, size=10, fill=PURPLE)
        set_cell_text(row.cells[1], value, size=10, fill=LIGHT if i % 2 == 0 else WHITE)
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(11.5)

    add_para(doc, "", space_after=4)
    add_para(
        doc,
        "Nota: material didáctico. Usar «No especificado» donde falte evidencia. "
        "No inventar credenciales ni datos reales.",
        size=9,
        color=MUTED,
    )
    doc.save(path)


# ---------------------------------------------------------------------------
# 3. Excel control
# ---------------------------------------------------------------------------
def build_control_xlsx(path: Path) -> None:
    wb = Workbook()

    thin = Border(
        left=Side(style="thin", color="D0CCE0"),
        right=Side(style="thin", color="D0CCE0"),
        top=Side(style="thin", color="D0CCE0"),
        bottom=Side(style="thin", color="D0CCE0"),
    )
    fill_header = PatternFill("solid", fgColor=PURPLE)
    fill_dark = PatternFill("solid", fgColor=DARK)
    fill_yellow = PatternFill("solid", fgColor=YELLOW)
    fill_light = PatternFill("solid", fgColor=LIGHT)
    fill_alt = PatternFill("solid", fgColor="EEEDF5")
    font_h = Font(name="Calibri", bold=True, color=WHITE, size=11)
    font_dark = Font(name="Calibri", color=DARK, size=10)
    font_bold = Font(name="Calibri", bold=True, color=DARK, size=10)
    font_title = Font(name="Calibri", bold=True, color=YELLOW, size=14)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def style_header(ws, headers, row=1):
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.fill = fill_header
            cell.font = font_h
            cell.alignment = center
            cell.border = thin

    def autosize(ws, widths):
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # --- Registro ---
    ws_r = wb.active
    ws_r.title = "Registro"
    ws_r["A1"] = "MCP-365-S2 · Registro de recepción · Proyecto Horizonte"
    ws_r["A1"].font = font_title
    ws_r["A1"].fill = fill_dark
    ws_r.merge_cells("A1:F1")
    ws_r.row_dimensions[1].height = 28

    reg_headers = [
        "Fecha correo",
        "Remitente",
        "Asunto",
        "Carpeta",
        "Adjunto",
        "Estado",
    ]
    style_header(ws_r, reg_headers, row=3)
    reg_row = [
        "2026-08-03",
        "Dirección de Operaciones",
        "NUEVO PROYECTO · HORIZONTE",
        "Proyectos/Horizonte/Entrada",
        "MCP365_S2_Ficha_Proyecto_Horizonte.docx",
        "Recibido · Pendiente de extracción",
    ]
    for col, val in enumerate(reg_row, 1):
        cell = ws_r.cell(row=4, column=col, value=val)
        cell.font = font_dark
        cell.alignment = left
        cell.border = thin
        cell.fill = fill_light
    # filas vacías para práctica
    for r in range(5, 10):
        for c in range(1, 7):
            cell = ws_r.cell(row=r, column=c, value="")
            cell.border = thin
            cell.fill = fill_alt if r % 2 == 0 else PatternFill()
    autosize(ws_r, [14, 28, 32, 28, 42, 30])
    ws_r["A11"] = "Nota: caso ficticio. Sin credenciales reales."
    ws_r["A11"].font = Font(name="Calibri", italic=True, color=MUTED, size=9)

    # --- Presupuesto ---
    ws_p = wb.create_sheet("Presupuesto")
    ws_p["A1"] = "MCP-365-S2 · Control presupuestario · Proyecto Horizonte"
    ws_p["A1"].font = font_title
    ws_p["A1"].fill = fill_dark
    ws_p.merge_cells("A1:K1")
    ws_p.row_dimensions[1].height = 28

    ws_p["A2"] = "Presupuesto preliminar (ficha)"
    ws_p["A2"].font = font_bold
    ws_p["B2"] = 100  # u.m.
    ws_p["B2"].font = Font(name="Calibri", bold=True, color=DARK, size=12)
    ws_p["B2"].fill = fill_yellow
    ws_p["B2"].border = thin
    ws_p["C2"] = "u.m. · valor de referencia de la ficha"
    ws_p["C2"].font = Font(name="Calibri", italic=True, color=MUTED, size=9)

    p_headers = [
        "Código partida",
        "Categoría",
        "Descripción",
        "Cantidad",
        "Costo unitario",
        "Costo total",
        "Responsable",
        "Fuente",
        "Tipo de cifra",
        "Estado de validación",
        "Observación",
    ]
    style_header(ws_p, p_headers, row=4)

    # qty, unit — total = formula
    budget_rows = [
        (
            "HZ-01",
            "Personal",
            "Coordinación PMO piloto (horas internas)",
            40,
            0.8,
            "Oficina de PMO",
            "Ficha / estimación interna",
            "Estimado",
            "Pendiente",
            "Tarifa interna de práctica",
        ),
        (
            "HZ-02",
            "Licencias",
            "Licencias de automatización (sandbox)",
            2,
            12,
            "TI Corporativa",
            "Cotización simulada",
            "Confirmado",
            "Pendiente",
            "Ambiente de laboratorio",
        ),
        (
            "HZ-03",
            "Capacitación",
            "Sesión de transferencia MCP-365-S2",
            1,
            8,
            "No especificado",
            "Ficha Proyecto Horizonte",
            "Estimado",
            "Pendiente",
            "Responsable por definir",
        ),
        (
            "HZ-04",
            "Consultoría",
            "Acompañamiento metodológico (días)",
            3,
            5,
            "Dirección de Operaciones",
            "Estimación patrocinador",
            "Estimado",
            "Pendiente",
            "",
        ),
        (
            "HZ-05",
            "Infraestructura",
            "Espacio SharePoint / OneDrive de práctica",
            1,
            4,
            "TI Corporativa",
            "Catálogo interno",
            "Confirmado",
            "Pendiente",
            "",
        ),
        (
            "HZ-06",
            "Contingencia",
            "Reserva para desviaciones",
            1,
            10,
            "Oficina de PMO",
            "Política piloto",
            "Estimado",
            "Pendiente",
            "10% del techo preliminar",
        ),
        (
            "HZ-07",
            "Comunicaciones",
            "Material de difusión interna del piloto",
            1,
            3.5,
            "No especificado",
            "No especificado",
            "No especificado",
            "Pendiente",
            "Monto y dueño por validar",
        ),
        (
            "HZ-08",
            "Calidad",
            "Revisión de control de calidad del informe",
            1,
            2.5,
            "Oficina de PMO",
            "Estimación PMO",
            "Estimado",
            "Pendiente",
            "",
        ),
    ]

    first_data = 5
    for i, row in enumerate(budget_rows):
        r = first_data + i
        code, cat, desc, qty, unit, resp, fuente, tipo, estado, obs = row
        values = [code, cat, desc, qty, unit, None, resp, fuente, tipo, estado, obs]
        for c, val in enumerate(values, 1):
            cell = ws_p.cell(row=r, column=c, value=val)
            cell.font = font_dark
            cell.alignment = left if c in (2, 3, 7, 8, 11) else center
            cell.border = thin
            cell.fill = fill_light if i % 2 == 0 else PatternFill()
        # F = D * E
        ws_p.cell(row=r, column=6, value=f"=D{r}*E{r}")
        ws_p.cell(row=r, column=6).number_format = "0.00"
        ws_p.cell(row=r, column=4).number_format = "0.00"
        ws_p.cell(row=r, column=5).number_format = "0.00"

    last_data = first_data + len(budget_rows) - 1
    sub_row = last_data + 1
    tot_row = last_data + 2
    des_row = last_data + 3

    ws_p.cell(row=sub_row, column=5, value="Subtotal").font = font_bold
    ws_p.cell(row=sub_row, column=6, value=f"=SUM(F{first_data}:F{last_data})")
    ws_p.cell(row=sub_row, column=6).font = font_bold
    ws_p.cell(row=sub_row, column=6).fill = fill_light
    ws_p.cell(row=sub_row, column=6).number_format = "0.00"
    ws_p.cell(row=sub_row, column=6).border = thin

    ws_p.cell(row=tot_row, column=5, value="Total partidas").font = font_bold
    ws_p.cell(row=tot_row, column=6, value=f"=F{sub_row}")
    ws_p.cell(row=tot_row, column=6).font = Font(name="Calibri", bold=True, color=WHITE, size=11)
    ws_p.cell(row=tot_row, column=6).fill = fill_header
    ws_p.cell(row=tot_row, column=6).number_format = "0.00"
    ws_p.cell(row=tot_row, column=6).border = thin

    ws_p.cell(row=des_row, column=5, value="Desviación vs preliminar").font = font_bold
    ws_p.cell(row=des_row, column=6, value=f"=F{tot_row}-B2")
    ws_p.cell(row=des_row, column=6).font = font_bold
    ws_p.cell(row=des_row, column=6).fill = fill_yellow
    ws_p.cell(row=des_row, column=6).number_format = "0.00"
    ws_p.cell(row=des_row, column=6).border = thin

    ws_p.cell(row=des_row + 2, column=1, value="Leyenda Tipo de cifra: Confirmado / Estimado / No especificado")
    ws_p.cell(row=des_row + 2, column=1).font = Font(name="Calibri", italic=True, color=MUTED, size=9)
    ws_p.cell(
        row=des_row + 3,
        column=1,
        value="Estado de validación inicial: Pendiente (no validado). Caso ficticio · sin credenciales reales.",
    ).font = Font(name="Calibri", italic=True, color=MUTED, size=9)

    autosize(ws_p, [14, 14, 42, 10, 14, 12, 22, 24, 14, 18, 28])

    # --- Riesgos ---
    ws_k = wb.create_sheet("Riesgos")
    ws_k["A1"] = "MCP-365-S2 · Registro de riesgos · Proyecto Horizonte"
    ws_k["A1"].font = font_title
    ws_k["A1"].fill = fill_dark
    ws_k.merge_cells("A1:M1")
    ws_k.row_dimensions[1].height = 28

    risk_headers = [
        "Identificador",
        "Riesgo",
        "Causa",
        "Consecuencia",
        "Probabilidad",
        "Impacto",
        "Nivel",
        "Control existente",
        "Tratamiento",
        "Responsable",
        "Fecha de revisión",
        "Estado",
        "Fuente",
    ]
    style_header(ws_k, risk_headers, row=3)

    example_risks = [
        (
            "R-HZ-01",
            "Campos incompletos en la ficha retrasan el arranque",
            "Líder de proyecto y observaciones figuran como No especificado",
            "Retraso en asignación de roles y en la cadena de automatización",
            "Media",
            "Alto",
            "Alto",
            "Revisión PMO al recibir ficha",
            "Completar campos faltantes antes de generar informe final",
            "Oficina de PMO",
            "2026-08-12",
            "No validado",
            "Ficha · Riesgos iniciales (R1)",
        ),
        (
            "R-HZ-02",
            "Desviación presupuestaria por cifras estimadas",
            "Varias partidas marcadas como Estimado o No especificado",
            "Superación del techo de 100 u.m. o replanificación forzada",
            "Media",
            "Medio",
            "Medio",
            "Hoja Presupuesto con desviación vs preliminar",
            "Validar tipos de cifra y responsables en comité",
            "Dirección de Operaciones",
            "2026-08-15",
            "No validado",
            "Ficha · Riesgos iniciales (R2) / Presupuesto",
        ),
        (
            "R-HZ-03",
            "Confusión entre automatización y aprobación humana",
            "Expectativa incorrecta de que Copilot pueda aprobar",
            "Informe o PPT generados sin control humano",
            "Baja",
            "Alto",
            "Medio",
            "Guía Power Automate · reto de Approvals",
            "Bloquear generación de PPT hasta aprobación humana",
            "Oficina de PMO",
            "2026-08-10",
            "No validado",
            "Ficha · Riesgos iniciales (R3)",
        ),
    ]

    for i, row in enumerate(example_risks):
        r = 4 + i
        for c, val in enumerate(row, 1):
            cell = ws_k.cell(row=r, column=c, value=val)
            cell.font = font_dark
            cell.alignment = left
            cell.border = thin
            cell.fill = fill_light if i % 2 == 0 else PatternFill()

    # filas vacías para riesgos inferidos
    for r in range(7, 15):
        for c in range(1, 14):
            cell = ws_k.cell(row=r, column=c, value="")
            cell.border = thin
            if c == 1 and r == 7:
                cell.value = "(espacio para riesgos inferidos)"
                cell.font = Font(name="Calibri", italic=True, color=MUTED, size=9)
            cell.fill = fill_alt if r % 2 == 0 else PatternFill()

    ws_k["A16"] = (
        "Estado de los riesgos de ejemplo: No validado. "
        "Completar filas vacías con riesgos inferidos durante la sesión. Caso ficticio."
    )
    ws_k["A16"].font = Font(name="Calibri", italic=True, color=MUTED, size=9)

    autosize(ws_k, [12, 36, 32, 32, 12, 10, 10, 28, 32, 20, 14, 12, 28])

    wb.save(path)


# ---------------------------------------------------------------------------
# 4. Plantilla informe Word
# ---------------------------------------------------------------------------
def build_informe_template(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    add_banner(
        doc,
        "MCP365_S2_Informe_Proyecto_Horizonte",
        "Plantilla de informe · Completar solo con evidencia · Caso ficticio",
    )
    add_para(
        doc,
        "Usar esta plantilla en el reto de Word + Approvals. Las cajas vacías "
        "deben llenarse con datos trazables. Validación humana y Control de calidad "
        "permanecen vacíos hasta la revisión humana (nunca por Copilot).",
        size=10,
        color=MUTED,
        space_after=12,
    )

    sections = [
        ("1. Identificación", "Proyecto, código, fechas, autor del informe, versión…"),
        ("2. Resumen ejecutivo", "Síntesis breve del estado y hallazgos principales…"),
        ("3. Objetivo", "Objetivo del proyecto según ficha / fuentes…"),
        ("4. Alcance", "Incluye / excluye…"),
        ("5. Entregables", "Lista de entregables y estado…"),
        ("6. Cronograma general", "Hitos y fechas clave…"),
        ("7. Presupuesto", "Resumen de partidas, total y tipo de cifra…"),
        ("8. Desviaciones", "Desviación vs presupuesto preliminar y causas…"),
        ("9. Riesgos prioritarios", "Riesgos de mayor nivel y tratamiento…"),
        ("10. Dependencias", "Dependencias activas o críticas…"),
        ("11. Decisiones pendientes", "Decisiones que requieren sponsor / PMO…"),
        ("12. Recomendaciones", "Acciones recomendadas con responsable…"),
        ("13. Fuentes", "Correo, ficha, Excel, JSON, etc. (citar rutas)…"),
        ("14. Validación humana", ""),  # EMPTY intentionally
        ("15. Control de calidad", ""),  # EMPTY intentionally
    ]

    for title, hint in sections:
        add_para(doc, title, size=12, bold=True, color=PURPLE, space_before=4, space_after=4)
        if title.startswith("14.") or title.startswith("15."):
            # EMPTY boxes — no placeholder text inside
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            shade_cell(cell, LIGHT)
            cell.text = ""
            for _ in range(3):
                cell.add_paragraph()
            # borde visual vía párrafo de ayuda fuera
            add_para(
                doc,
                "← Sección reservada: debe permanecer vacía hasta validación humana. No completar con Copilot.",
                size=8,
                color=MUTED,
                space_after=8,
            )
        else:
            empty_box(doc, hint)

    add_para(
        doc,
        "MCP-365-S2 · Plantilla · Proyecto Horizonte (ficticio) · Sin credenciales reales",
        size=8,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.save(path)


# ---------------------------------------------------------------------------
# 5. Guía PDF Power Automate
# ---------------------------------------------------------------------------
def build_guia_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title="MCP365_S2_Guia_Power_Automate",
        author="MCP-365 · Caso ficticio",
    )
    styles = getSampleStyleSheet()
    c_dark = colors.HexColor(f"#{DARK}")
    c_purple = colors.HexColor(f"#{PURPLE}")
    c_yellow = colors.HexColor(f"#{YELLOW}")
    c_light = colors.HexColor(f"#{LIGHT}")
    c_muted = colors.HexColor(f"#{MUTED}")

    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            textColor=c_yellow,
            alignment=TA_CENTER,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSub",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=11,
            textColor=colors.white,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HReto",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=c_purple,
            spaceBefore=10,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyES",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            textColor=c_dark,
            alignment=TA_JUSTIFY,
            leading=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Warn",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=c_dark,
            backColor=c_yellow,
            borderPadding=6,
            leading=13,
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="FootNote",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            textColor=c_muted,
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            textColor=c_dark,
            leading=13,
        )
    )

    story = []

    # Cover banner
    cover = Table(
        [
            [
                Paragraph("MCP-365-S2 · Guía Power Automate", styles["CoverTitle"]),
            ],
            [
                Paragraph(
                    "Cadena de 6 retos · Proyecto Horizonte (caso ficticio)",
                    styles["CoverSub"],
                )
            ],
            [
                Paragraph(
                    "Sin CDN · Sin credenciales reales · Aprobación humana obligatoria",
                    styles["CoverSub"],
                )
            ],
        ],
        colWidths=[16 * cm],
    )
    cover.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), c_dark),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ]
        )
    )
    story.append(cover)
    story.append(Spacer(1, 0.4 * cm))

    story.append(
        Paragraph(
            "Esta guía describe la cadena de automatización didáctica para la sesión "
            "MCP-365-S2. El objetivo es conectar correo, adjunto, JSON, Excel, Word con "
            "Approvals y PowerPoint <b>solo después</b> de la aprobación humana.",
            styles["BodyES"],
        )
    )
    story.append(
        Paragraph(
            "REGLA CRÍTICA: la aprobación humana NUNCA la realiza Copilot ni un modelo de IA. "
            "Approvals de Power Automate debe esperar la decisión de una persona autorizada "
            "(PMO / patrocinador). Copilot puede redactar borradores; no puede firmar ni aprobar.",
            styles["Warn"],
        )
    )

    retos = [
        (
            "Reto 1 — Disparador de correo + filtro de asunto",
            [
                "Crear un flujo con disparador «Al llegar un nuevo correo» (Outlook).",
                "Filtrar por asunto exacto: <b>NUEVO PROYECTO · HORIZONTE</b>.",
                "Verificar remitente de práctica: Dirección de Operaciones.",
                "Si el asunto no coincide, no iniciar la cadena.",
            ],
        ),
        (
            "Reto 2 — Guardar adjunto",
            [
                "Obtener el adjunto <b>MCP365_S2_Ficha_Proyecto_Horizonte.docx</b>.",
                "Guardarlo en la carpeta del proyecto (p. ej. Proyectos/Horizonte/Entrada).",
                "Registrar en la hoja <b>Registro</b> del Excel de control: fecha, remitente, asunto, carpeta, adjunto, estado.",
            ],
        ),
        (
            "Reto 3 — Extraer JSON estructurado",
            [
                "A partir de la ficha, producir un JSON con campos: nombre, código, objetivo, alcance, fechas, presupuesto, entregables, riesgos, etc.",
                "Donde falte evidencia, usar literalmente <b>No especificado</b> (no inventar).",
                "Persistir el JSON en la carpeta del proyecto para trazabilidad.",
            ],
        ),
        (
            "Reto 4 — Excel: presupuesto y riesgos",
            [
                "Actualizar la hoja <b>Presupuesto</b>: partidas, tipo de cifra (Confirmado/Estimado/No especificado), fórmulas de total y desviación.",
                "Actualizar la hoja <b>Riesgos</b> con riesgos de la ficha y espacio para riesgos inferidos.",
                "Dejar Estado de validación / Estado del riesgo como no validados hasta revisión humana.",
            ],
        ),
        (
            "Reto 5 — Word + Approvals (humano)",
            [
                "Rellenar la plantilla <b>MCP365_S2_Plantilla_Informe_Proyecto.docx</b> con evidencia.",
                "Dejar vacías las secciones <b>Validación humana</b> y <b>Control de calidad</b>.",
                "Enviar Approvals a la persona autorizada. Esperar Aprobado / Rechazado.",
                "<b>Prohibido:</b> simular aprobación con Copilot o auto-aprobar el flujo.",
            ],
        ),
        (
            "Reto 6 — PowerPoint solo tras aprobación",
            [
                "Condición: si Approvals = Aprobado → generar PPT de cierre de fase.",
                "Si Rechazado → notificar, no generar PPT, registrar motivo.",
                "El PPT resume: identificación, presupuesto, riesgos, decisiones y próximas acciones.",
            ],
        ),
    ]

    for title, bullets in retos:
        story.append(Paragraph(title, styles["HReto"]))
        items = [
            ListItem(Paragraph(b, styles["BulletBody"]), leftIndent=8, value="•")
            for b in bullets
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="•"))

    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Orden de la cadena", styles["HReto"]))
    chain = (
        "Correo (filtro asunto) → Guardar adjunto → Extraer JSON → "
        "Excel (presupuesto/riesgos) → Informe Word + Approvals humano → "
        "PPT (solo si aprobado)"
    )
    chain_tbl = Table(
        [[Paragraph(chain, styles["BulletBody"])]],
        colWidths=[16 * cm],
    )
    chain_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), c_light),
                ("BOX", (0, 0), (-1, -1), 1.5, c_purple),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(chain_tbl)
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph(
            "MCP-365-S2 · Proyecto Horizonte · Material didáctico · "
            f"Paleta #{DARK} · #{PURPLE} · #{YELLOW}",
            styles["FootNote"],
        )
    )

    doc.build(story)


# ---------------------------------------------------------------------------
# 6. Resultado esperado HTML
# ---------------------------------------------------------------------------
def build_resultado_html(path: Path) -> None:
    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>MCP365_S2 · Resultado esperado · Proyecto Horizonte</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: Calibri, Segoe UI, Arial, sans-serif;
    background: #E8E6F0;
    color: #{DARK};
    line-height: 1.55;
    padding: 24px 16px 40px;
  }}
  .wrap {{ max-width: 900px; margin: 0 auto; }}
  header {{
    background: #{DARK};
    color: #fff;
    padding: 22px 26px;
    border-bottom: 4px solid #{YELLOW};
  }}
  header .tag {{
    color: #{YELLOW};
    font-size: 12px;
    letter-spacing: 0.06em;
    text-transform: uppercase;
    margin-bottom: 6px;
  }}
  header h1 {{ font-size: 22px; margin-bottom: 8px; }}
  header p {{ font-size: 13px; color: #C8C4D8; }}
  .banner {{
    background: #{YELLOW};
    color: #{DARK};
    font-size: 13px;
    font-weight: 700;
    padding: 10px 16px;
    margin-bottom: 18px;
  }}
  section {{
    background: #fff;
    border: 1px solid #D0CCE0;
    margin-bottom: 14px;
    padding: 16px 18px;
  }}
  section h2 {{
    color: #{PURPLE};
    font-size: 15px;
    margin-bottom: 10px;
    padding-bottom: 6px;
    border-bottom: 2px solid #{LIGHT};
  }}
  .grid {{
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 10px;
  }}
  @media (max-width: 700px) {{
    .grid {{ grid-template-columns: 1fr; }}
  }}
  .card {{
    background: #{LIGHT};
    padding: 10px 12px;
    font-size: 13px;
  }}
  .card strong {{ display: block; color: #{PURPLE}; margin-bottom: 4px; }}
  ul {{ margin: 6px 0 0 18px; font-size: 13px; }}
  li {{ margin-bottom: 4px; }}
  pre {{
    background: #{DARK};
    color: #E8E6F0;
    padding: 12px 14px;
    font-size: 12px;
    overflow-x: auto;
    line-height: 1.4;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 13px;
    margin-top: 6px;
  }}
  th {{
    background: #{PURPLE};
    color: #fff;
    text-align: left;
    padding: 8px 10px;
  }}
  td {{
    border: 1px solid #D0CCE0;
    padding: 7px 10px;
  }}
  tr:nth-child(even) td {{ background: #{LIGHT}; }}
  .ok {{ color: #{PURPLE}; font-weight: 700; }}
  footer {{
    text-align: center;
    font-size: 11px;
    color: #{MUTED};
    margin-top: 10px;
  }}
  code {{
    background: #{LIGHT};
    padding: 1px 5px;
    font-size: 12px;
    color: #{PURPLE};
  }}
</style>
</head>
<body>
  <div class="wrap">
    <header>
      <div class="tag">MCP-365-S2 · Referencia de calidad</div>
      <h1>Resultado esperado · Proyecto Horizonte</h1>
      <p>
        Vista de referencia del resultado extremo a extremo de la sesión.
        No es fuente de análisis: sirve para contrastar la calidad del entregable.
        Caso ficticio · sin credenciales reales.
      </p>
    </header>
    <div class="banner">
      ⚠ Referencia de calidad · No usar como fuente de análisis ni para inventar datos faltantes
    </div>

    <section>
      <h2>1. Carpeta del proyecto (estructura esperada)</h2>
      <div class="grid">
        <div class="card">
          <strong>Entrada</strong>
          Correo simulado + ficha Word guardada<br>
          <code>…/Horizonte/Entrada/</code>
        </div>
        <div class="card">
          <strong>Control</strong>
          Excel presupuesto/riesgos + JSON extraído<br>
          <code>…/Horizonte/Control/</code>
        </div>
        <div class="card">
          <strong>Informes</strong>
          Informe Word + evidencia de Approvals<br>
          <code>…/Horizonte/Informes/</code>
        </div>
        <div class="card">
          <strong>Presentaciones</strong>
          PPT generado solo tras aprobación<br>
          <code>…/Horizonte/Presentaciones/</code>
        </div>
      </div>
    </section>

    <section>
      <h2>2. Muestra JSON (extracción)</h2>
      <pre>{{
  "proyecto": "Proyecto Horizonte",
  "codigo": "PRJ-HZ-2026-01",
  "objetivo": "Habilitar flujo piloto de gobernanza documental y control presupuestario…",
  "patrocinador": "Dirección de Operaciones",
  "lider": "No especificado",
  "fecha_inicio": "2026-08-10",
  "fecha_fin": "2026-11-30",
  "presupuesto_preliminar_um": 100,
  "observaciones": "No especificado",
  "riesgos_iniciales": ["R1", "R2", "R3"],
  "fuente": "MCP365_S2_Ficha_Proyecto_Horizonte.docx"
}}</pre>
    </section>

    <section>
      <h2>3. Resumen de presupuesto (referencia)</h2>
      <table>
        <thead>
          <tr>
            <th>Concepto</th>
            <th>Valor</th>
            <th>Nota</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td>Presupuesto preliminar (ficha)</td>
            <td>100 u.m.</td>
            <td>Celda de referencia en hoja Presupuesto</td>
          </tr>
          <tr>
            <td>Partidas de ejemplo</td>
            <td>8 (HZ-01 … HZ-08)</td>
            <td>Mezcla Confirmado / Estimado / No especificado</td>
          </tr>
          <tr>
            <td>Total partidas</td>
            <td>Fórmula =SUM(…)</td>
            <td>Calculado en Excel, no hardcodear en el informe</td>
          </tr>
          <tr>
            <td>Desviación</td>
            <td>Total − preliminar</td>
            <td>Explicar causas; validación humana pendiente</td>
          </tr>
        </tbody>
      </table>
    </section>

    <section>
      <h2>4. Riesgos (referencia)</h2>
      <ul>
        <li><strong>R-HZ-01</strong> — Campos incompletos (líder No especificado) · Estado: No validado · Fuente: ficha</li>
        <li><strong>R-HZ-02</strong> — Desviación por cifras estimadas · Estado: No validado · Fuente: ficha / presupuesto</li>
        <li><strong>R-HZ-03</strong> — Confusión automatización vs aprobación humana · Estado: No validado · Fuente: ficha</li>
        <li>Filas vacías reservadas para riesgos <em>inferidos</em> durante la sesión</li>
      </ul>
    </section>

    <section>
      <h2>5. Informe aprobado (criterio de aceptación)</h2>
      <ul>
        <li>Plantilla completada con evidencia; sin inventar campos faltantes</li>
        <li>Secciones <strong>Validación humana</strong> y <strong>Control de calidad</strong> llenadas solo por persona autorizada</li>
        <li>Approvals en estado <span class="ok">Aprobado</span> (humano, nunca Copilot)</li>
        <li>Fuentes citadas: correo, ficha, Excel, JSON</li>
      </ul>
    </section>

    <section>
      <h2>6. Outline PPT (solo post-aprobación)</h2>
      <ol style="margin-left:18px;font-size:13px;">
        <li>Portada · Proyecto Horizonte · PRJ-HZ-2026-01</li>
        <li>Objetivo y alcance</li>
        <li>Presupuesto y desviación</li>
        <li>Riesgos prioritarios</li>
        <li>Decisiones y próximas acciones</li>
        <li>Evidencia de aprobación humana</li>
      </ol>
    </section>

    <footer>
      MCP-365-S2 · Resultado esperado · Paleta #{DARK} · #{PURPLE} · #{YELLOW}<br>
      Referencia de calidad · no fuente de análisis · caso ficticio
    </footer>
  </div>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    targets = {
        "MCP365_S2_Correo_Inicio_Proyecto_Horizonte.html": build_email_html,
        "MCP365_S2_Ficha_Proyecto_Horizonte.docx": build_ficha_docx,
        "MCP365_S2_Control_Proyecto_Horizonte.xlsx": build_control_xlsx,
        "MCP365_S2_Plantilla_Informe_Proyecto.docx": build_informe_template,
        "MCP365_S2_Guia_Power_Automate.pdf": build_guia_pdf,
        "MCP365_S2_Resultado_Esperado.html": build_resultado_html,
    }

    print(f"Salida: {OUT_DIR}")
    for name, builder in targets.items():
        path = OUT_DIR / name
        builder(path)
        size = path.stat().st_size
        print(f"  OK  {name}  ({size:,} bytes)")

    print("Listo.")


if __name__ == "__main__":
    main()
