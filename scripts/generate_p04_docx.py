# -*- coding: utf-8 -*-
"""Genera fuente WhatsApp ST-14 y plantilla oficial MCP-365-P04 · Tres audiencias."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_FUENTE = ROOT / "planillas" / "MCP365_P04_Fuente_chat_WhatsApp_ST14_SU7.docx"
OUT_PLANILLA = ROOT / "planillas" / "MCP365_P04_Comunicacion_tres_audiencias.docx"

PURPLE = "4A3DA6"
DARK = "221E40"
YELLOW = "FFEC00"
LIGHT = "F4F3F9"
HINT_BG = "FFF8CC"
RULE_BG = "EFEAFB"
CHAT_BG = "E8F5E9"

MESSAGES = [
    ("6:47 a. m.", "Andrés Molina", "Buenos días. Inicio turno de operación de 06:00 a 14:00 en la Subestación Urbana 7. Los equipos principales están operando sin alarmas activas y el servicio se encuentra normal."),
    ("7:12 a. m.", "Andrés Molina", "Durante la inspección rutinaria encontré una lectura intermitente en el sensor de temperatura ST-14 del tablero de control secundario."),
    ("7:13 a. m.", "Andrés Molina", "La pantalla alterna entre una lectura de 42 °C y el mensaje “sin señal”. La pérdida de lectura dura entre 6 y 12 segundos y luego se recupera."),
    ("7:15 a. m.", "Andrés Molina", "Aclaro que no se ha presentado disparo, alarma de sobretemperatura ni interrupción del servicio."),
    ("7:16 a. m.", "Laura Méndez", "Andrés, activa el protocolo de verificación PV-OPS-07. Por favor, no concluyamos todavía que existe calentamiento real hasta comparar con una medición independiente."),
    ("7:18 a. m.", "Andrés Molina", "Recibido. Inicio PV-OPS-07 y registro el hallazgo como ST-14-SU7-180326."),
    ("7:24 a. m.", "Andrés Molina", "Primera medición manual con termómetro infrarrojo: 43,1 °C en el punto asociado al ST-14."),
    ("7:25 a. m.", "Andrés Molina", "Los sensores cercanos están estables: ST-13 registra 41,8 °C y ST-15 registra 42,6 °C."),
    ("7:27 a. m.", "Andrés Molina", "No observo olor, humo, decoloración, vibración anormal ni evidencia visible de recalentamiento en el tablero."),
    ("7:30 a. m.", "Felipe Castro", "Con esos datos no tenemos evidencia de una condición térmica anormal. La falla podría estar en el sensor, en el conector o en la transmisión de la señal, pero todavía es una hipótesis."),
    ("7:32 a. m.", "Felipe Castro", "No registremos “sensor dañado” como causa confirmada. La causa raíz solo podrá definirse después de revisar conexiones y probar el sensor retirado."),
    ("7:35 a. m.", "Laura Méndez", "De acuerdo. Estado del hallazgo: lectura intermitente confirmada; calentamiento real no evidenciado; causa raíz pendiente."),
    ("7:39 a. m.", "Andrés Molina", "Segunda medición manual: 43,4 °C. La lectura se mantiene dentro del rango observado en el tablero."),
    ("7:42 a. m.", "Andrés Molina", "Revisé visualmente el cable y el conector accesible sin intervenir el circuito. No se observan daños externos ni conexiones sueltas a simple vista."),
    ("7:45 a. m.", "Felipe Castro", "Para revisar terminales internamente debemos aislar el circuito secundario y aplicar el procedimiento de seguridad. Esa actividad no se debe realizar durante la operación rutinaria sin una ventana autorizada."),
    ("7:47 a. m.", "Laura Méndez", "Entonces mantengamos el equipo en servicio con seguimiento reforzado. Andrés, registra una medición manual cada dos horas y en cada cambio de turno."),
    ("7:49 a. m.", "Andrés Molina", "Confirmado. Haré seguimiento a las 09:30, 11:30 y 13:30 y dejaré el registro para el siguiente turno."),
    ("7:52 a. m.", "Jorge Salazar", "¿Existe afectación actual para usuarios o riesgo inmediato de interrupción?"),
    ("7:54 a. m.", "Laura Méndez", "No hay afectación actual a usuarios. El servicio continúa normal y no se ha activado ninguna protección."),
    ("7:55 a. m.", "Laura Méndez", "Con la información disponible, el impacto se clasifica como bajo, siempre que mantengamos la vigilancia y no aparezca una temperatura real anormal."),
    ("7:58 a. m.", "Felipe Castro", "Propongo los siguientes criterios de escalamiento:\n1. Temperatura manual igual o superior a 55 °C.\n2. Pérdida continua de señal durante más de 10 minutos.\n3. Diferencia superior a 8 °C frente a los sensores cercanos.\n4. Activación de alarma, olor, humo o cambio visible en el tablero."),
    ("8:01 a. m.", "Jorge Salazar", "Aprobados esos criterios para el seguimiento temporal. Si se cumple cualquiera, se debe informar inmediatamente a Operaciones y evaluar adelantar la intervención."),
    ("8:05 a. m.", "Laura Méndez", "Abrí la orden de trabajo OT-MNT-2026-0417 para revisión del circuito, reemplazo preventivo del ST-14 y prueba funcional."),
    ("8:08 a. m.", "Laura Méndez", "La intervención queda programada para el domingo 5 de abril de 2026, entre las 09:00 y las 11:00."),
    ("8:10 a. m.", "Jorge Salazar", "¿La intervención requiere suspender el servicio?"),
    ("8:12 a. m.", "Felipe Castro", "No se tiene prevista una interrupción del servicio. El trabajo se realizará sobre el circuito secundario del sensor, con aislamiento del punto de intervención."),
    ("8:14 a. m.", "Felipe Castro", "La ventana de dos horas incluye preparación, aislamiento, reemplazo, verificación de conexiones, prueba funcional y observación de estabilidad."),
    ("8:17 a. m.", "Felipe Castro", "El reemplazo físico debería tomar aproximadamente 25 minutos. Después necesitamos 15 minutos de señal estable y comparación con medición manual."),
    ("8:20 a. m.", "Laura Méndez", "La ejecución requiere técnico de instrumentación, técnico de operación, elementos de protección personal y autorización de ingreso al área."),
    ("8:23 a. m.", "Paula Rincón", "Confirmo que en el almacén central existe una unidad compatible con el ST-14."),
    ("8:24 a. m.", "Paula Rincón", "Código de inventario: INV-ST14-A. Es la única unidad disponible actualmente."),
    ("8:26 a. m.", "Paula Rincón", "Puedo reservarla desde hoy para la orden OT-MNT-2026-0417 y programar su despacho para el viernes 3 de abril antes de las 2:00 p. m."),
    ("8:28 a. m.", "Laura Méndez", "Paula, por favor deja la unidad reservada. No se requiere compra adicional para esta intervención."),
    ("8:29 a. m.", "Paula Rincón", "Reserva confirmada. Responsable de entrega: Almacén Central. Fecha prevista: 3 de abril de 2026, antes de las 2:00 p. m."),
    ("8:33 a. m.", "Jorge Salazar", "Mantengamos la fecha del 5 de abril mientras las mediciones continúen estables. No considero necesaria una intervención de emergencia con la evidencia actual."),
    ("8:35 a. m.", "Jorge Salazar", "Realizaremos una validación de continuidad o “go/no-go” el sábado 4 de abril a las 4:00 p. m. Deben estar confirmados el repuesto, el personal y la autorización de ingreso."),
    ("8:38 a. m.", "Laura Méndez", "Queda registrado. Yo consolidaré esas tres confirmaciones y reportaré el resultado de la validación."),
    ("8:42 a. m.", "Camila Torres", "Para comunicaciones: ¿debemos informar desde ahora a la comunidad aledaña?"),
    ("8:44 a. m.", "Jorge Salazar", "No. En este momento no existe interrupción ni afectación a usuarios. No se autoriza todavía publicar un aviso."),
    ("8:46 a. m.", "Laura Méndez", "Durante la intervención se restringirá únicamente el acceso inmediato al tablero. No se prevé cerrar vías, suspender el servicio ni impedir el acceso general al sector."),
    ("8:48 a. m.", "Felipe Castro", "La restricción directa alrededor del tablero se estima en 20 a 25 minutos. El resto de la ventana corresponde a pruebas y observación."),
    ("8:51 a. m.", "Camila Torres", "Entendido. Prepararé un borrador preventivo, pero no lo publicaré hasta recibir instrucción de Operaciones."),
    ("8:53 a. m.", "Camila Torres", "¿En qué situación debemos activar la comunicación?"),
    ("8:55 a. m.", "Jorge Salazar", "Si el cierre del área o la restricción de acceso se extiende por más de 30 minutos, Comunicaciones debe emitir el aviso preventivo."),
    ("8:57 a. m.", "Jorge Salazar", "También se debe informar si cambia la fecha, si aparece una afectación al servicio o si la intervención requiere una restricción adicional para la comunidad."),
    ("9:00 a. m.", "Camila Torres", "Confirmado. El aviso incluirá únicamente fecha, franja horaria, posible restricción de acceso, ausencia de corte programado y canal de contacto."),
    ("9:02 a. m.", "Camila Torres", "El canal autorizado para este caso será la Línea Comunitaria 01 8000 000 707, opción 2. Este número es ficticio y se utiliza únicamente para el ejercicio."),
    ("9:30 a. m.", "Andrés Molina", "Seguimiento de las 09:30: medición manual de 43,2 °C. El sensor continúa presentando pérdidas breves de señal, pero no hay alarma ni cambio en la operación."),
    ("9:34 a. m.", "Felipe Castro", "Gracias. Mantener la observación. La intermitencia continúa, pero no hay evidencia nueva que cambie la clasificación de impacto bajo."),
    ("10:05 a. m.", "Laura Méndez", "Resumen provisional del caso:\n* Hallazgo: lectura intermitente del ST-14.\n* Ubicación: tablero de control secundario de la Subestación Urbana 7.\n* Servicio: normal, sin afectación a usuarios.\n* Protocolo activado: PV-OPS-07.\n* Causa raíz: pendiente de confirmación.\n* Impacto actual: bajo.\n* Repuesto: reservado en Almacén Central.\n* Intervención: 5 de abril, de 09:00 a 11:00.\n* Comunicación externa: no autorizada por ahora."),
    ("10:09 a. m.", "Jorge Salazar", "El resumen es correcto. Para gerencia debemos destacar que no existe afectación, no se requiere compra adicional y el riesgo permanece controlado con monitoreo."),
    ("10:12 a. m.", "Jorge Salazar", "También debe quedar claro que la decisión vigente es ejecutar el reemplazo preventivo en la ventana programada, salvo que se active un criterio de escalamiento."),
    ("10:16 a. m.", "Felipe Castro", "Para el informe técnico incluyan la hipótesis como no confirmada. No afirmar que el sensor está dañado hasta finalizar las pruebas."),
    ("10:20 a. m.", "Camila Torres", "Para la comunidad debemos evitar códigos como PV-OPS-07, INV-ST14-A y OT-MNT-2026-0417. No son necesarios para explicar la situación."),
    ("10:22 a. m.", "Camila Torres", "Tampoco se debe mencionar “falla térmica”, porque las mediciones manuales no muestran sobretemperatura. Podemos hablar de mantenimiento preventivo de un dispositivo de monitoreo."),
    ("11:30 a. m.", "Andrés Molina", "Seguimiento de las 11:30: 43,3 °C. Sin alarmas, sin interrupción y sin cambios visibles."),
    ("11:34 a. m.", "Laura Méndez", "Recibido. Continúa el seguimiento hasta entregar el turno."),
    ("1:30 p. m.", "Andrés Molina", "Último seguimiento del turno: 43,1 °C. El servicio permanece normal. Las pérdidas de señal siguen siendo breves e intermitentes."),
    ("1:36 p. m.", "Andrés Molina", "Dejo fotografías, mediciones y registro de eventos asociados a la orden OT-MNT-2026-0417 para consulta del siguiente turno."),
    ("1:40 p. m.", "Laura Méndez", "Gracias. Cierre del reporte del turno:\n* No hubo interrupción del servicio.\n* No se presentó afectación a usuarios.\n* No se confirmó sobretemperatura.\n* El ST-14 quedó bajo monitoreo reforzado.\n* El reemplazo preventivo continúa programado para el 5 de abril.\n* La causa raíz sigue pendiente de confirmación.\n* El aviso a la comunidad solo se emitirá si se cumple alguno de los criterios definidos."),
]


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def set_cell_text(cell, text, *, bold=False, color=None, size=10, fill=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or DARK)
    if fill:
        shade_cell(cell, fill)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "60")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=9, fill=PURPLE)
    for r_i, row_data in enumerate(rows):
        fill = LIGHT if r_i % 2 == 1 else None
        for c_i, val in enumerate(row_data):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, size=9, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_callout(doc, title, body, bg):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    set_run(p.add_run(title), size=10, bold=True, color=DARK)
    set_run(p.add_run(" " + body), size=10, color=DARK)
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), size=13, bold=True, color=PURPLE)


def add_empty_box(doc, hint: str):
    add_para(doc, hint, size=9, color="5A5A72", space_after=4)
    table = doc.add_table(rows=4, cols=1)
    for row in table.rows:
        set_cell_text(row.cells[0], "", size=10, fill=LIGHT)
    doc.add_paragraph()


def setup_page(doc):
    for s in doc.sections:
        s.top_margin = Cm(1.4)
        s.bottom_margin = Cm(1.4)
        s.left_margin = Cm(1.6)
        s.right_margin = Cm(1.6)


def add_banner(doc, title: str):
    banner = doc.add_table(rows=1, cols=1)
    cell = banner.rows[0].cells[0]
    shade_cell(cell, DARK)
    cell.text = ""
    set_run(cell.paragraphs[0].add_run(title), size=15, bold=True, color=YELLOW)
    p2 = cell.add_paragraph()
    set_run(p2.add_run("Misión Copilot 365 · ECCO · Universidad Sergio Arboleda"), size=9, color="D5D2E6")
    doc.add_paragraph()


def add_chips(doc, left: str, right: str):
    chips = doc.add_table(rows=1, cols=2)
    set_cell_text(chips.rows[0].cells[0], left, bold=True, color="FFFFFF", size=9, fill=PURPLE, center=True)
    set_cell_text(chips.rows[0].cells[1], right, bold=True, color=DARK, size=9, fill=YELLOW, center=True)
    doc.add_paragraph()


def add_message(doc, time: str, author: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, CHAT_BG)
    cell.text = ""
    head = cell.paragraphs[0]
    set_run(head.add_run(f"{time} · {author}"), size=10, bold=True, color=PURPLE)
    for line in body.split("\n"):
        p = cell.add_paragraph()
        set_run(p.add_run(line if line else " "), size=10, color=DARK)
    doc.add_paragraph()


def build_fuente():
    doc = Document()
    setup_page(doc)
    add_banner(doc, "Fuente · Chat de WhatsApp · Hallazgo ST-14")
    add_chips(doc, "MCP-365-P04-F", "Solo lectura · Caso ficticio")
    add_callout(
        doc,
        "Uso:",
        "Este archivo es la FUENTE. No es la plantilla de llenado. "
        "Ábrelo junto con MCP365_P04_Comunicacion_tres_audiencias.docx. "
        "Copilot analiza este chat y completa solo la plantilla de tres audiencias.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Advertencia:",
        "Caso ficticio y anonimizado. El número de la Línea Comunitaria es ficticio y solo se usa para el ejercicio.",
        HINT_BG,
    )

    heading(doc, "Datos del grupo")
    add_table(
        doc,
        ["Campo", "Información"],
        [
            ["Grupo", "Operación y Mantenimiento · Subestación Urbana 7"],
            ["Fecha", "miércoles, 18 de marzo de 2026"],
            ["Caso", "Hallazgo ST-14"],
            ["Participantes", "Andrés Molina · Laura Méndez · Felipe Castro · Paula Rincón · Camila Torres · Jorge Salazar"],
        ],
        col_widths=[4, 12],
    )

    heading(doc, "Transcripción del chat (no editar)")
    for time, author, body in MESSAGES:
        add_message(doc, time, author, body)

    OUT_FUENTE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FUENTE)
    print(f"OK -> {OUT_FUENTE}")


def build_planilla():
    doc = Document()
    setup_page(doc)
    add_banner(doc, "Plantilla de llenado · Comunicación a tres audiencias")
    add_chips(doc, "MCP-365-P04", "Word + Copilot · Llenar")

    add_table(
        doc,
        ["Campo", "Dato"],
        [
            ["Participante", "________________________________"],
            ["Área / rol", "________________________________"],
            ["Fecha", "________________"],
            ["App usada", "Word + Copilot"],
            ["Fuente usada", "MCP365_P04_Fuente_chat_WhatsApp_ST14_SU7.docx"],
            ["Estado", "☐ Borrador    ☐ Validado    ☐ Entregado"],
        ],
        col_widths=[4, 12],
    )
    doc.add_paragraph()

    add_callout(
        doc,
        "Archivos del reto:",
        "1) Fuente: MCP365_P04_Fuente_chat_WhatsApp_ST14_SU7.docx · "
        "2) Esta plantilla. Completa secciones 1–6. Validación humana, CQ y firmas: solo la persona.",
        RULE_BG,
    )
    add_callout(
        doc,
        "Entrega:",
        "Guarda como MCP365_P04_Tres_audiencias_completado.docx. Conserva tablas y diseño. "
        "No inventes hechos. No afirmes causa raíz confirmada ni “falla térmica”.",
        HINT_BG,
    )

    heading(doc, "1. Fuente técnica (extracto clave) · LLENAR CON COPILOT")
    add_empty_box(doc, "Resume hechos verificables del chat (hallazgo, mediciones, protocolo, impacto, intervención, comunicación).")

    heading(doc, "2. Versión A · Resumen técnico · LLENAR CON COPILOT")
    add_empty_box(doc, "Audiencia: especialistas. Prioriza causa/hipótesis no confirmada, protocolo, acción, controles y criterios de escalamiento. Lenguaje especializado permitido.")

    heading(doc, "3. Versión B · Resumen para gerencia · LLENAR CON COPILOT")
    add_empty_box(doc, "Audiencia: dirección. Prioriza impacto, riesgo, decisión vigente, plazo y recursos. Máximo 10 líneas.")

    heading(doc, "4. Versión C · Comunicación a comunidad / usuarios · LLENAR CON COPILOT")
    add_empty_box(
        doc,
        "Audiencia general. Qué ocurre, cuándo, cómo afecta y a quién contactar. Sin códigos técnicos (PV-OPS-07, INV-ST14-A, OT-MNT…). Sin “falla térmica”. Aviso solo si corresponde; no inventar publicación ya emitida.",
    )

    heading(doc, "5. Matriz de consistencia de hechos · LLENAR CON COPILOT")
    empty5 = [["", "☐ Sí ☐ No", "☐ Sí ☐ No", "☐ Sí ☐ No", "☐ Sí ☐ No", ""] for _ in range(5)]
    add_table(
        doc,
        ["Hecho original", "¿En versión técnica?", "¿En gerencia?", "¿En comunidad?", "¿Se alteró el significado?", "Corrección"],
        empty5,
        col_widths=[4, 2.2, 2.2, 2.2, 2.6, 2.6],
    )

    heading(doc, "6. Información omitida a propósito · LLENAR CON COPILOT")
    add_table(
        doc,
        ["Versión", "Qué se omitió", "Por qué", "¿Riesgo de malentendido?"],
        [
            ["Técnica", "", "", "☐ Bajo ☐ Medio ☐ Alto"],
            ["Gerencia", "", "", "☐ Bajo ☐ Medio ☐ Alto"],
            ["Comunidad", "", "", "☐ Bajo ☐ Medio ☐ Alto"],
        ],
        col_widths=[3, 5, 5, 3],
    )

    heading(doc, "7. Control de calidad · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Criterio", "Cumple", "Observaciones"],
        [
            ["Hechos solo del chat WhatsApp", "☐ Sí ☐ No", ""],
            ["Causa raíz / sensor dañado no afirmados como confirmados", "☐ Sí ☐ No", ""],
            ["Versión comunidad sin códigos técnicos innecesarios", "☐ Sí ☐ No", ""],
            ["Diseño de plantilla conservado", "☐ Sí ☐ No", ""],
        ],
        col_widths=[7, 3, 6],
    )

    heading(doc, "8. Firmas · SOLO LA PERSONA (no Copilot)")
    add_table(
        doc,
        ["Rol", "Nombre", "Fecha"],
        [["Elaboró", "", ""], ["Revisó", "", ""], ["Aprobó", "", ""]],
        col_widths=[4, 8, 4],
    )

    add_para(
        doc,
        "Misión Copilot 365 · Formato oficial MCP-365-P04 · No reconstruir el diseño.",
        size=8,
        color="5A5A72",
        space_after=0,
    )

    doc.save(OUT_PLANILLA)
    print(f"OK -> {OUT_PLANILLA}")


if __name__ == "__main__":
    build_fuente()
    build_planilla()
